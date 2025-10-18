# tender_analyzer.py - Script to perform tender analysis and generate documents using NumPy-based vector search

import google.generativeai as genai
import os
import pypdf
import json
import re
import yaml
import pandas as pd
import numpy as np
from typing import List, Dict, Any, Tuple, Optional
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- Configuration ---
API_KEY = "AIzaSyBWKV0tJU89iMcRq3ak9-ssNYRAiRnlj2s"
MODEL_NAME = "gemini-flash-latest"
EMBEDDING_MODEL_NAME = "models/text-embedding-004"

# --- Global variables ---
CONFIG = {}
VECTOR_DB: Optional[np.ndarray] = None
VECTOR_DB_METADATA: List[Dict[str, Any]] = []
EMBEDDING_SIZE = 768

config_path = os.path.join(
    os.path.dirname(__file__), "..", "master.yml"
)
config_path = os.path.abspath(config_path)
#config_path = r"C:\Users\Nandita\OneDrive\Desktop\Tryingtender\TenderResponse\master.yml"
# ---------------- CONFIGURATION ----------------
def load_config(config_path):
    """Loads configuration from the master.yml file."""
    global CONFIG
    try:
        with open(config_path, 'r', encoding='utf-8') as f:
            CONFIG = yaml.safe_load(f)
        if 'working_directory' in CONFIG and os.path.isdir(CONFIG['working_directory']):
            os.chdir(CONFIG['working_directory'])
        else:
            print(f"WARNING: 'working_directory' not found or invalid in {config_path}. Using current directory.")
    except FileNotFoundError:
        raise FileNotFoundError(f"Configuration file '{config_path}' not found. Please create it.")
    except yaml.YAMLError as e:
        raise yaml.YAMLError(f"Error parsing YAML file '{config_path}': {e}")


# ---------------- VECTOR DB (NUMPY) ----------------
def load_vector_db():
    """Loads precomputed embeddings and metadata stored as NumPy arrays and JSON."""
    global VECTOR_DB, VECTOR_DB_METADATA
    try:
        load_config(config_path)
        output_path = os.path.join(os.getcwd(), CONFIG['output_folder'])
        embeddings_path = os.path.join(output_path, 'vector_db.npy')
        metadata_path = os.path.join(output_path, 'vector_db_metadata.json')

        if not os.path.exists(embeddings_path) or not os.path.exists(metadata_path):
            raise FileNotFoundError("Vector database files not found. Please run rag_builder.py first.")

        VECTOR_DB = np.load(embeddings_path)
        with open(metadata_path, 'r', encoding='utf-8') as f:
            VECTOR_DB_METADATA = json.load(f)

        print(f"✅ Loaded NumPy vector database with {len(VECTOR_DB_METADATA)} documents.")
    except Exception as e:
        print(f"ERROR: Failed to load vector database. {e}")
        VECTOR_DB = None
        VECTOR_DB_METADATA = []


# ---------------- MODEL SETUP ----------------
def get_gemini_model():
    """Initializes and returns a GenerativeModel instance."""
    if not API_KEY:
        raise ValueError("GEMINI_API_KEY not found. Please set it or provide it directly.")
    genai.configure(api_key=API_KEY)
    return genai.GenerativeModel(MODEL_NAME)


def get_embedding_model() -> GoogleGenerativeAIEmbeddings:
    """Initializes and returns a LangChain GoogleGenerativeAIEmbeddings instance."""
    if not API_KEY:
        raise ValueError("GEMINI_API_KEY not found. Please set it or provide it directly.")
    return GoogleGenerativeAIEmbeddings(model=EMBEDDING_MODEL_NAME, google_api_key=API_KEY)


# ---------------- EMBEDDING ----------------
def embed_text(text: str) -> List[float]:
    """Generates an embedding for the given text using LangChain."""
    try:
        model = get_embedding_model()
        embedding = model.embed_query(text)
        return embedding
    except Exception as e:
        print(f"ERROR: Failed to generate embedding. Exception: {e}")
        return []


# ---------------- PDF EXTRACTION ----------------
def extract_text_from_pdf(pdf_path: str) -> Optional[str]:
    """Extracts text from a PDF document."""
    text = ""
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file not found at {pdf_path}")
        return None
    try:
        with open(pdf_path, 'rb', encoding='utf-8') as file:
            reader = pypdf.PdfReader(file)
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text()
    except Exception as e:
        print(f"Error reading PDF {pdf_path}: {e}")
        return None
    return text


# ---------------- GEMINI ANALYSIS ----------------
def analyze_document_with_gemini_text_output(prompt: str, document_text: str, context: str = "") -> str:
    """Sends document text and context to Gemini API with a specific prompt, expecting plain text output."""
    try:
        model = get_gemini_model()
        print("model",model)

        # Prevent overly large payloads (>50k chars)
        if len(document_text) > 50000:
            document_text = document_text[:50000]
        if len(context) > 30000:
            context = context[:30000]
        print("if complete")

        full_prompt = f"{prompt}\n\nRelevant Context:\n{context}\n\nDocument:\n{document_text}"
        print("full prom")

        response = model.generate_content(full_prompt)
        print("respon",response)

        # Gemini responses have 'text' attribute in newer SDKs, but sometimes nested in candidates
        if hasattr(response, "text"):
            return response.text.strip()
        elif hasattr(response, "candidates") and response.candidates:
            return response.candidates[0].content.parts[0].text.strip()
        else:
            return "⚠️ No valid response text received from Gemini."

    except Exception as e:
        print(f"\n🔥 Gemini API Error: {e}\n")
        return f"Analysis failed due to API error: {e}"

# def analyze_document_with_gemini_text_output(prompt: str, document_text: str, context: str = "") -> str:
#     """Sends document text and context to Gemini API with a specific prompt, expecting plain text output."""
#     try:
#         model = get_gemini_model()
#         full_prompt = f"{prompt}\n\nRelevant Context:\n{context}\n\nDocument:\n{document_text}"
#         response = model.generate_content([full_prompt])
#         return response.text
#     except Exception as e:
#         print(f"ERROR: Failed to call Gemini API. Exception: {e}")
#         return "Analysis failed due to API error."


# ---------------- DOCUMENT RETRIEVAL (NUMPY) ----------------
def cosine_similarity(a: np.ndarray, b: np.ndarray) -> float:
    """Computes cosine similarity between two vectors."""
    if np.linalg.norm(a) == 0 or np.linalg.norm(b) == 0:
        return 0.0
    return np.dot(a, b) / (np.linalg.norm(a) * np.linalg.norm(b))


def retrieve_relevant_documents(query_text: str, top_k: int = 5) -> List[Dict[str, Any]]:
    """Retrieves most relevant documents using cosine similarity."""
    if VECTOR_DB is None or not VECTOR_DB_METADATA:
        print("WARNING: Vector database not initialized or empty. Cannot retrieve documents.")
        return []

    query_embedding = np.array(embed_text(query_text))
    if query_embedding.size == 0:
        return []

    similarities = np.array([cosine_similarity(query_embedding, vec) for vec in VECTOR_DB])
    top_indices = similarities.argsort()[-top_k:][::-1]

    return [VECTOR_DB_METADATA[i] for i in top_indices]


# ---------------- CONTEXT FORMATTING ----------------
def format_retrieved_context(retrieved_docs: List[Dict[str, Any]]) -> str:
    """Formats retrieved documents into a string for LLM context."""
    context_str = ""
    for i, doc in enumerate(retrieved_docs):
        context_str += f"--- Retrieved Document {i+1} (Type: {doc['type']}, Source: {doc['source']}) ---\n"
        context_str += doc['text'] + "\n\n"
    return context_str


# ---------------- ANALYSIS STEPS ----------------
def get_bid_decision_and_reasoning_with_rag(tender_text: str, company_profile_text: str, retrieved_context: str):
    """Uses Gemini with RAG to make a bid/no-bid decision and provide reasoning."""
    prompt = (
        "Analyze the Tender Document and Company Profile. Use the provided context to decide "
        "whether the company should bid ('yes') or not ('no'). Then provide a short explanation. "
        "Response format: yes/no: explanation.\n\n"
    )
    full_input = f"Tender Document:\n{tender_text}\n\nCompany Profile:\n{company_profile_text}"
    response = analyze_document_with_gemini_text_output(prompt, full_input, context=retrieved_context)

    decision, explanation = "no", "Could not parse response."
    match = re.match(r"^(yes|no):\s*(.*)", response.strip().lower(), re.DOTALL)
    if match:
        decision, explanation = match.groups()
    elif "yes" in response.lower():
        decision = "yes"
        explanation = "Inferred yes, but format mismatch."
    elif "no" in response.lower():
        decision = "no"
        explanation = "Inferred no, but format mismatch."

    explanation = explanation.capitalize()
    return decision, explanation


def get_tender_response_outline_with_rag(tender_text: str, retrieved_context: str="") -> str:
    """Generates a structured outline for tender response."""
    prompt = (
        "Based on the Tender Document and context, create a structured tender response outline "
        "with sections like Executive Summary, Technical Bid, Financial Bid, and Annexures."
    )
    return analyze_document_with_gemini_text_output(prompt, tender_text, context=retrieved_context)
def get_estimated_cost_table_with_rag(tender_text: str, retrieved_context: str = "") -> List[Dict[str, Any]]:
    """
    Provides structured cost estimates as a LIST of dicts.
    Expected format returned by model (JSON array):
      [
        {"component": "Material", "min_inr": 10000, "max_inr": 15000, "notes": "..."}, ...
      ]
    Fallback: returns one auto-generated component estimate.
    """
    prompt = (
        "Analyze the Tender Document and context to provide estimated cost breakdowns in INR. "
        "Return output strictly as a JSON array with objects having keys: component, min_inr, max_inr, notes."
    )
    try:
        response = analyze_document_with_gemini_text_output(prompt, tender_text, context=retrieved_context)
        # try to find the JSON array inside the response
        json_match = re.search(r"(\[.*\])", response, re.DOTALL)
        if not json_match:
            # maybe the whole response is JSON
            try:
                parsed = json.loads(response)
            except Exception:
                parsed = None
        else:
            parsed = json.loads(json_match.group(1))

        # Validate parsed structure
        if isinstance(parsed, list) and parsed:
            cleaned = []
            for item in parsed:
                if not isinstance(item, dict):
                    continue
                comp = item.get("component", "Component")
                try:
                    min_inr = float(item.get("min_inr", 0))
                except Exception:
                    min_inr = 0.0
                try:
                    max_inr = float(item.get("max_inr", 0))
                except Exception:
                    max_inr = min_inr
                notes = str(item.get("notes", ""))
                cleaned.append({
                    "component": comp,
                    "min_inr": round(min_inr, 2),
                    "max_inr": round(max_inr, 2),
                    "notes": notes
                })
            if cleaned:
                return cleaned

        # If parsing or validation failed, raise to trigger fallback
        raise ValueError("Parsed response is not a valid list of dicts.")
    except Exception as e:
        # fallback heuristic (make a single-line estimate)
        base = max(1000.0, len(tender_text) * 1.5)  # crude heuristic
        low = round(base * 0.8, 2)
        high = round(base * 1.3, 2)
        return [{"component": "AutoEstimate", "min_inr": low, "max_inr": high, "notes": f"Fallback estimate ({e})"}]

# def get_estimated_cost_table_with_rag(tender_text: str, retrieved_context: str = "") -> Tuple[float, float]:
#     """Provides structured cost estimates as JSON."""
#     prompt = (
#         "Analyze the Tender Document and context to provide estimated cost breakdowns in INR. "
#         "Return output strictly as JSON array with keys: component, min_inr, max_inr, notes."
#     )
#     try:
#         estimates = analyze_document_with_gemini_text_output(prompt, tender_text, context=retrieved_context)  # returns list of dicts
#         if not estimates:
#             raise ValueError("Empty estimates, using fallback.")
#         total_min = sum(item["min_inr"] for item in estimates)
#         total_max = sum(item["max_inr"] for item in estimates)
#         return total_min, total_max
#     except Exception:
#         # fallback heuristic
#         base = len(tender_text) * 1.5
#         return round(base * 0.8, 2), round(base * 1.3, 2)

def get_win_probability_with_rag(tender_text: str, retrieved_context: str = "") -> float:
    """
    Estimates the probability of winning the tender based on tender details and context.

    Args:
        tender_text (str): The full tender document text.
        retrieved_context (str, optional): Retrieved RAG context for more accurate evaluation.

    Returns:
        float: Win probability (between 0.0 and 1.0).
    """
    prompt = (
        "Analyze the tender document and the provided organizational context. "
        "Estimate the probability (in percentage) that the company will win this tender. "
        "Consider factors such as eligibility, technical capabilities, pricing, "
        "previous experience, and competition. Return only a single numeric value "
        "between 0 and 100 representing the win probability percentage."
    )

    try:
        response = analyze_document_with_gemini_text_output(prompt, tender_text, context=retrieved_context)
        # Extract numeric probability safely
        import re
        match = re.search(r"(\d+(\.\d+)?)", response)
        if match:
            probability = float(match.group(1)) / 100  # Convert to 0–1 scale
            return min(max(probability, 0.0), 1.0)
        else:
            return 0.0
    except Exception as e:
        print(f"Error calculating win probability: {e}")
        return 0.0

def get_competitor_win_probabilities_with_rag(
    tender_text: str,
    competitors: list,
    retrieved_context: str = ""
) -> dict:
    """
    Given a list of competitors, estimate their win probability for the same tender.
    Returns a dict: {competitor_name: probability (0–1 scale)}
    """
    results = {}

    if not competitors:
        return {"No competitors found": 0.0}

    for comp in competitors:
        prompt = (
            f"You are analyzing a tender. Based on the tender details and context, "
            f"estimate the probability (in percentage) that the competitor '{comp}' "
            f"will win the tender. Consider eligibility, technical strength, pricing, "
            f"past experience, and market competition. Return only a single numeric "
            f"value between 0 and 100."
        )
        try:
            response = analyze_document_with_gemini_text_output(prompt, tender_text, context=retrieved_context)
            import re
            match = re.search(r"(\d+(\.\d+)?)", response)
            if match:
                prob = float(match.group(1)) / 100
                results[comp] = round(min(max(prob, 0.0), 1.0), 2)
            else:
                results[comp] = 0.0
        except Exception as e:
            print(f"Error estimating probability for {comp}: {e}")
            results[comp] = 0.0

    return results

def get_competitor_list_with_rag(tender_text: str, retrieved_context: str = "") -> list:
    """
    Extracts or predicts potential competitors from the tender and contextual documents.

    Returns:
        list: A list of competitor company names (strings)
    """
    prompt = (
        "Analyze the following tender document and contextual information. "
        "Identify potential or mentioned competitors who might bid for this tender. "
        "If competitors are not directly mentioned, predict likely competitors "
        "based on industry, project type, and scale. "
        "Return the competitor names as a numbered list, without explanations."
    )

    try:
        response = analyze_document_with_gemini_text_output(prompt, tender_text, context=retrieved_context)
        # Extract clean list of competitor names
        competitors = []
        for line in response.splitlines():
            line = line.strip("-•1234567890. \t")
            if line and len(line.split()) < 10:  # avoid sentences
                competitors.append(line)
        # Remove duplicates and empty lines
        competitors = list(dict.fromkeys([c for c in competitors if c]))
        return competitors or ["No competitors identified."]
    except Exception as e:
        print(f"Error generating competitor list: {e}")
        return ["Error generating competitor list."]

# def get_estimated_cost_table_with_rag(tender_text: str, retrieved_context: str="") -> List[Dict[str, Any]]:
#     """Provides structured cost estimates as JSON."""
#     prompt = (
#         "Analyze the Tender Document and context to provide estimated cost breakdowns in INR. "
#         "Return output strictly as JSON array with keys: component, min_inr, max_inr, notes."
#     )
#     response = analyze_document_with_gemini_text_output(prompt, tender_text, context=retrieved_context)
#     try:
#         json_match = re.search(r'(\[.*\])', response, re.DOTALL)
#         if json_match:
#             parsed = json.loads(json_match.group(1))
#             if isinstance(parsed, list):
#                 return parsed
#     except Exception:
#         pass
#     return []


# ---------------- SAVE OUTPUTS ----------------
def save_analysis_outputs_to_docx(tender_name: str, analysis_data: Dict[str, Any]):
    """Saves results to DOCX and Excel."""
    try:
        load_config(config_path)
        output_dir = os.path.join(os.getcwd(), CONFIG['output_folder'], tender_name)
        os.makedirs(output_dir, exist_ok=True)

        # --- Bid Decision DOCX ---
        doc_bid = Document()
        doc_bid.add_heading(f"Tender Analysis: {tender_name}", level=1)
        doc_bid.add_heading("Bid Decision:", level=2)
        doc_bid.add_paragraph(analysis_data['final_bid_decision'].upper())
        doc_bid.add_heading("Reasoning:", level=2)
        doc_bid.add_paragraph(analysis_data['final_decision_reasoning'])
        doc_bid.save(os.path.join(output_dir, f"{tender_name}_BidDecision.docx"))

        # --- Outline DOCX (if applicable) ---
        if analysis_data['final_bid_decision'].lower() == "yes":
            doc_outline = Document()
            doc_outline.add_heading(f"Tender Response Outline: {tender_name}", level=1)
            for line in analysis_data['tender_response_outline'].split('\n'):
                if line.strip().startswith('#'):
                    doc_outline.add_heading(line.strip('# ').strip(), level=2)
                elif line.strip():
                    doc_outline.add_paragraph(line.strip())
            doc_outline.save(os.path.join(output_dir, f"{tender_name}_TenderResponseOutline.docx"))

        # --- Cost Estimates Excel ---
        if analysis_data['cost_estimates_data']:
            df = pd.DataFrame(analysis_data['cost_estimates_data'])
            df.loc[len(df.index)] = {
                "component": "Total Estimated Cost",
                "min_inr": df['min_inr'].sum(),
                "max_inr": df['max_inr'].sum(),
                "notes": "Sum of all components"
            }
            excel_path = os.path.join(output_dir, f"{tender_name}_EstimatedCost.xlsx")
            df.to_excel(excel_path, index=False)
            print(f"✅ Saved cost estimates to: {excel_path}")

    except Exception as e:
        print(f"Error saving documents: {e}")


# ---------------- MAIN PIPELINE ----------------
def run_analysis():
    """Main function to run the tender analysis pipeline."""
    try:
        load_config(config_path)
        load_vector_db()

        tender_dir = os.path.join(os.getcwd(), CONFIG['tender_documents_folder'])
        company_dir = os.path.join(os.getcwd(), CONFIG['company_documents_folder'])

        company_profile_path = os.path.join(company_dir, CONFIG['company_profile_filename'])
        company_turnover_path = os.path.join(company_dir, CONFIG['company_turnover_filename'])

        tender_files = [f for f in os.listdir(tender_dir) if f.endswith('.pdf')]
        if not tender_files:
            print(f"No tender documents found in {tender_dir}.")
            return

        for tender_file in tender_files:
            tender_name = os.path.splitext(tender_file)[0]
            tender_path = os.path.join(tender_dir, tender_file)

            tender_text = extract_text_from_pdf(tender_path)
            company_profile = extract_text_from_pdf(company_profile_path)
            company_turnover = extract_text_from_pdf(company_turnover_path)

            if not tender_text or not company_profile:
                print(f"Skipping {tender_name}: Missing document text.")
                continue

            query_text = f"{tender_text}\n\n{company_profile}\n\n{company_turnover}"
            retrieved_docs = retrieve_relevant_documents(query_text)
            retrieved_context = format_retrieved_context(retrieved_docs)

            final_decision, reasoning = get_bid_decision_and_reasoning_with_rag(
                tender_text, company_profile + "\n" + company_turnover, retrieved_context
            )

            outline = ""
            if final_decision.lower() == "yes":
                outline = get_tender_response_outline_with_rag(tender_text, retrieved_context)

            cost_estimates = get_estimated_cost_table_with_rag(tender_text, retrieved_context)

            analysis_data = {
                "final_bid_decision": final_decision,
                "final_decision_reasoning": reasoning,
                "tender_response_outline": outline,
                "cost_estimates_data": cost_estimates
            }

            save_analysis_outputs_to_docx(tender_name, analysis_data)
            print(f"✅ Completed analysis for: {tender_name}")

    except Exception as e:
        print(f"An error occurred during analysis: {e}")


if __name__ == "__main__":
    run_analysis()
